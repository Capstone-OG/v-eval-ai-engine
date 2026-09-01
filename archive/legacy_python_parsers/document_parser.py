import os
import json
import tempfile
from PIL import Image
from pypdf import PdfReader
import docx

def parse_pdf(file_path):
    """
    Trích xuất text từ file PDF sử dụng pypdf.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Không tìm thấy file: {file_path}")
        
    reader = PdfReader(file_path)
    pages_text = []
    
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages_text.append({
                "page_number": i + 1,
                "text": text.strip()
            })
            
    return {
        "format": "PDF",
        "file_name": os.path.basename(file_path),
        "total_pages": len(reader.pages),
        "pages": pages_text
    }

def parse_docx(file_path):
    """
    Trích xuất text từ file DOCX sử dụng python-docx.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Không tìm thấy file: {file_path}")
        
    doc = docx.Document(file_path)
    paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                "paragraph_index": i + 1,
                "text": para.text.strip()
            })
            
    # Hỗ trợ trích xuất dữ liệu từ bảng biểu (Tables) nếu có
    tables_data = []
    for t_idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)
        tables_data.append({
            "table_index": t_idx + 1,
            "rows": rows_data
        })
        
    return {
        "format": "DOCX",
        "file_name": os.path.basename(file_path),
        "paragraphs_count": len(paragraphs),
        "paragraphs": paragraphs,
        "tables": tables_data
    }

def parse_doc(file_path):
    """
    Trích xuất text từ file DOC (đời cũ) bằng cách gọi COM Automation (win32com) trên Windows.
    Yêu cầu máy phải có cài đặt Microsoft Word.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Không tìm thấy file: {file_path}")
        
    # Thử import win32com để gọi MS Word
    try:
        import win32com.client
    except ImportError:
        return {
            "format": "DOC",
            "file_name": os.path.basename(file_path),
            "status": "Error",
            "error_message": "Chưa cài đặt thư viện 'pywin32' để hỗ trợ COM. Chạy: pip install pywin32"
        }
        
    word = None
    temp_docx_path = None
    try:
        # Khởi tạo Word Application (chạy ngầm)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Mở file .doc
        abs_doc_path = os.path.abspath(file_path)
        doc = word.Documents.Open(abs_doc_path)
        
        # Lưu thành file .docx tạm thời
        temp_dir = tempfile.gettempdir()
        temp_docx_path = os.path.join(temp_dir, os.path.basename(file_path) + "x")
        
        # 16 đại diện cho định dạng wdFormatXMLDocument (.docx)
        doc.SaveAs2(temp_docx_path, FileFormat=16)
        doc.Close()
        
        # Parse file .docx vừa chuyển đổi
        result = parse_docx(temp_docx_path)
        result["format"] = "DOC (Converted to DOCX)"
        return result
        
    except Exception as e:
        return {
            "format": "DOC",
            "file_name": os.path.basename(file_path),
            "status": "Error",
            "error_message": f"Lỗi COM Automation: {str(e)}. Hãy chắc chắn Microsoft Word được cài đặt."
        }
    finally:
        if word:
            try:
                word.Quit()
            except:
                pass
        if temp_docx_path and os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except:
                pass

def parse_image(file_path):
    """
    Đọc siêu dữ liệu hình ảnh bằng Pillow (PIL).
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Không tìm thấy file: {file_path}")
        
    with Image.open(file_path) as img:
        return {
            "format": "IMAGE",
            "file_name": os.path.basename(file_path),
            "image_format": img.format,
            "width": img.width,
            "height": img.height,
            "mode": img.mode, # e.g. RGB, RGBA
            "info": str(img.info) # EXIF/Metadata khác nếu có
        }
